"""
Module 1: Document Processor
Handles PDF, DOCX, TXT and other text-based document formats.
Extracts text, page info, and paragraph structure.
"""

import json
import re
import uuid
from pathlib import Path
from typing import Any

# PDF processing
try:
    import pymupdf4llm
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# DOCX processing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

DATA_DIR = Path("data")
UPLOADS_DIR = DATA_DIR / "uploads"
PROCESSED_DIR = DATA_DIR / "processed"


def ensure_dirs():
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)


def generate_doc_id() -> str:
    return str(uuid.uuid4())[:8]


def _extract_paragraphs(text: str) -> list[str]:
    """Split text into meaningful paragraphs."""
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    # Further split very long single-line blocks
    result = []
    for p in paragraphs:
        if len(p) > 2000:
            # Split by sentence
            sentences = re.split(r'(?<=[.!?])\s+', p)
            chunk, chunks = "", []
            for s in sentences:
                if len(chunk) + len(s) < 800:
                    chunk += (" " if chunk else "") + s
                else:
                    if chunk:
                        chunks.append(chunk)
                    chunk = s
            if chunk:
                chunks.append(chunk)
            result.extend(chunks)
        else:
            result.append(p)
    return result


def process_pdf(file_path: Path) -> dict[str, Any]:
    """Extract text and structure from a PDF file."""
    ensure_dirs()

    if not PYMUPDF_AVAILABLE:
        raise RuntimeError("PyMuPDF not installed. Run: uv add pymupdf pymupdf4llm")

    pages = []
    doc = fitz.open(str(file_path))
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        paragraphs = _extract_paragraphs(text)
        pages.append({
            "page": page_num + 1,
            "text": text.strip(),
            "paragraphs": paragraphs,
            "word_count": len(text.split()),
            "char_count": len(text),
        })
    doc.close()

    # Also get markdown representation via pymupdf4llm
    try:
        md_text = pymupdf4llm.to_markdown(str(file_path))
    except Exception:
        md_text = ""

    return {
        "document_name": file_path.name,
        "format": "pdf",
        "total_pages": len(pages),
        "pages": pages,
        "markdown": md_text,
        "full_text": "\n\n".join(p["text"] for p in pages),
    }


def process_docx(file_path: Path) -> dict[str, Any]:
    """Extract text and structure from a DOCX file."""
    ensure_dirs()

    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Run: uv add python-docx")

    document = docx.Document(str(file_path))
    all_paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    # Group paragraphs into "sections" based on heading styles
    sections = []
    current_section = {"section": 1, "heading": "Document Start", "paragraphs": []}
    section_num = 1

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name:
            if current_section["paragraphs"]:
                sections.append(current_section)
            section_num += 1
            current_section = {
                "section": section_num,
                "heading": text,
                "paragraphs": [],
            }
        else:
            current_section["paragraphs"].append(text)

    if current_section["paragraphs"] or not sections:
        sections.append(current_section)

    # Convert sections to page-like structure for uniformity
    pages = []
    for i, section in enumerate(sections):
        text = section["heading"] + "\n\n" + "\n\n".join(section["paragraphs"])
        pages.append({
            "page": i + 1,
            "heading": section["heading"],
            "text": text.strip(),
            "paragraphs": section["paragraphs"],
            "word_count": len(text.split()),
            "char_count": len(text),
        })

    full_text = "\n\n".join(p["text"] for p in pages)
    return {
        "document_name": file_path.name,
        "format": "docx",
        "total_pages": len(pages),
        "pages": pages,
        "markdown": full_text,
        "full_text": full_text,
    }


def process_txt(file_path: Path) -> dict[str, Any]:
    """Extract text from a plain text file."""
    ensure_dirs()

    text = file_path.read_text(encoding="utf-8", errors="replace")
    paragraphs = _extract_paragraphs(text)

    # Chunk into pages of ~500 words each
    pages = []
    current_page_words = []
    current_page_paras = []
    page_num = 1

    for para in paragraphs:
        current_page_paras.append(para)
        current_page_words.extend(para.split())
        if len(current_page_words) >= 500:
            page_text = "\n\n".join(current_page_paras)
            pages.append({
                "page": page_num,
                "text": page_text,
                "paragraphs": current_page_paras,
                "word_count": len(current_page_words),
                "char_count": len(page_text),
            })
            page_num += 1
            current_page_words = []
            current_page_paras = []

    if current_page_paras:
        page_text = "\n\n".join(current_page_paras)
        pages.append({
            "page": page_num,
            "text": page_text,
            "paragraphs": current_page_paras,
            "word_count": len(current_page_words),
            "char_count": len(page_text),
        })

    if not pages:
        pages = [{"page": 1, "text": text, "paragraphs": paragraphs,
                  "word_count": len(text.split()), "char_count": len(text)}]

    return {
        "document_name": file_path.name,
        "format": "txt",
        "total_pages": len(pages),
        "pages": pages,
        "markdown": text,
        "full_text": text,
    }


def process_document(file_path: Path, doc_id: str) -> dict[str, Any]:
    """Route document to appropriate processor based on file extension."""
    ensure_dirs()
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        result = process_pdf(file_path)
    elif ext == ".docx":
        result = process_docx(file_path)
    elif ext in (".txt", ".md", ".rst", ".csv"):
        result = process_txt(file_path)
    else:
        # Attempt generic text extraction
        result = process_txt(file_path)

    result["doc_id"] = doc_id
    result["file_path"] = str(file_path)

    # Save to data/processed/
    out_path = PROCESSED_DIR / f"{doc_id}.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    return result


def load_processed_document(doc_id: str) -> dict[str, Any] | None:
    """Load a previously processed document."""
    path = PROCESSED_DIR / f"{doc_id}.json"
    if not path.exists():
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def list_processed_documents() -> list[dict[str, Any]]:
    """List all processed documents with summary metadata."""
    ensure_dirs()
    docs = []
    for p in PROCESSED_DIR.glob("*.json"):
        try:
            with open(p, "r", encoding="utf-8") as f:
                data = json.load(f)
            docs.append({
                "doc_id": data.get("doc_id", p.stem),
                "document_name": data.get("document_name", p.name),
                "format": data.get("format", "unknown"),
                "total_pages": data.get("total_pages", 0),
                "file_path": data.get("file_path", ""),
            })
        except Exception:
            pass
    return docs
